
import re
import json
import datetime
from io import BytesIO
from pathlib import Path
import streamlit as st
from docx import Document

st.set_page_config(page_title="Prompts Gu Smart V4", page_icon="🧠", layout="wide")

PLACEHOLDER_RE = re.compile(r"\[\s*([A-Z0-9_]+)\s*\]")
BASE_MARKER = "[COLE AQUI O BLOCO BASE GLOBAL]"

DEFAULTS = {
    "NOME_DO_ALUNO": "Gustavo",
    "APELIDO": "Gu",
    "IDADE": "9 anos",
    "ANO_SERIE": "4º ano",
    "ESCOLA": "Colégio Albert Sabin",
    "TURNO_DO_ALUNO": "manhã",
    "INTERESSES_DO_ALUNO": "Minecraft, desafios, exemplos visuais, jogos rápidos",
    "NOME_DO_RESPONSAVEL": "Vanessa",
    "NIVEL_DO_ALUNO": "aprende melhor com exemplos visuais, precisa de progressão curta e foco"
}

BASE_BLOCK_TEMPLATE = """PROFESSOR PARTICULAR DE {NOME_DO_ALUNO}

Você é o professor particular de {NOME_DO_ALUNO}.
Você é especialista em pedagogia do Ensino Fundamental I, neuroaprendizagem, adaptação para TDAH e apoio a dificuldades de processamento auditivo.

PERFIL DO ALUNO
Nome: {NOME_DO_ALUNO}
Apelido: {APELIDO}
Idade: {IDADE}
Série/Ano: {ANO_SERIE}
Escola: {ESCOLA}
Turno: {TURNO_DO_ALUNO}

COMO O ALUNO APRENDE
• Atenção curta: priorize foco, clareza e progressão curta
• Dificuldade de processamento auditivo: valorize organização visual, exemplos concretos e linguagem objetiva
• Aprende melhor com exemplos e interação
• Precisa ganhar confiança, mas também ser desafiado na medida certa

REGRA PEDAGÓGICA PRINCIPAL
Explique de forma clara e acessível, mas SEM infantilizar o tom, SEM simplificar demais o raciocínio e SEM tratar o aluno como incapaz.

NÍVEL DE DESAFIO
Toda produção deve ter progressão em 3 camadas:
1. compreensão básica
2. aplicação
3. desafio leve ou transferência

USO DAS FONTES
• Use as fontes apenas para entender conceitos, habilidades cobradas, vocabulário e estilo de cobrança
• NÃO copiar exercícios, frases ou exemplos do livro
• Criar exemplos inéditos
• Quando a fonte trouxer um exemplo, manter apenas a habilidade e trocar contexto, números, objetos e pergunta
• Evitar depender apenas dos exemplos do livro

VARIEDADE DE CONTEXTOS
• Não usar sempre os mesmos temas
• Pode usar interesses do aluno, mas com moderação
• Alternar entre cotidiano, escola, dinheiro, jogos, esporte, comida, coleções, tempo, medidas, natureza, tecnologia e desafios lógicos
• Se usar os interesses do aluno, usar como apoio pontual e não como base de tudo

INTERESSES DO ALUNO
{INTERESSES_DO_ALUNO}

TOM
• Claro
• Respeitoso
• Encorajador
• Objetivo
• Intelectualmente honesto
• Sem “voz de bebê”, sem exagero de fofura, sem excesso de elogios vazios
"""

CRONOGRAMA_TEMPLATE = """{BASE}

PROMPT — CRONOGRAMA COMPLETO ATÉ A PROVA

CONTEXTO ATUAL
Data de hoje: {DATA_DE_HOJE}
Data da prova: {DATA_DA_PROVA}
Atenção para outros eventos no mesmo período: {OUTRAS_PROVAS_NO_PERIODO}
Matéria: {MATERIA}

CONTEÚDOS DA PROVA
{CONTEUDOS_DA_PROVA}

PRIORIDADE
Alta:
{PRIORIDADES}

Média:
{CONTEUDOS_MEDIOS}

Baixa:
{CONTEUDOS_SECUNDARIOS}

REGRAS DE PLANEJAMENTO
• Dividir o estudo por dias até a prova
• Sessões de 15 a 25 minutos
• Máximo de 1 conteúdo principal por dia
• Priorizar primeiro os conteúdos de maior dificuldade e maior chance de cair
• Incluir revisão final obrigatória no dia anterior à prova
• Não sobrecarregar
• Se houver pouco tempo, reduzir conteúdos secundários
• O último dia antes da prova deve focar revisão e consolidação
• No dia da prova, não incluir estudo formal; apenas descanso ou revisão mental leve

IMPORTANTE
• NÃO explicar conteúdo
• NÃO ensinar
• NÃO gerar exemplos
• NÃO detalhar a aula
• NÃO dividir o dia em dois blocos
• Gerar o plano completo, nunca apenas um dia

FORMATO OBRIGATÓRIO

[DIA X]
Conteúdo do dia:
Duração:
Objetivo:
"""

def safe_format(template, values):
    data = {}
    data.update(DEFAULTS)
    data.update(values)
    return template.format(**data)

def recommend_material(days_left, situacao, prioridade):
    if days_left <= 1:
        return "revisão estratégica + exercícios curtos + mini quiz + orientação ao responsável"
    if situacao == "novo":
        return "vídeo explicativo curto + exemplos progressivos + prática guiada + mini quiz"
    if situacao == "em_dificuldade":
        return "explicação enxuta + exercícios guiados passo a passo + correção comentada + mini quiz"
    if prioridade == "alta":
        return "explicação curta + prática guiada + desafio leve + revisão de erros"
    return "resumo rápido + aplicação + mini quiz"

def recommend_intensity(days_left, situacao):
    if days_left <= 1:
        return "baixa a moderada, foco em segurança"
    if situacao == "em_dificuldade":
        return "moderada com progressão curta"
    if situacao == "novo":
        return "moderada com construção gradual"
    return "moderada com treino"

def recommend_mode(days_left, situacao):
    if days_left <= 1:
        return "pré-prova"
    if situacao == "novo":
        return "introdução guiada"
    if situacao == "em_dificuldade":
        return "retomada estratégica"
    return "treino com consolidação"

def make_prompt(kind, values):
    common = safe_format("""{BASE}

CONTEXTO
Matéria: {MATERIA}
Conteúdo do dia: {CONTEUDO_DO_DIA}
Data da prova: {DATA_DA_PROVA}
Dias restantes: {DIAS_RESTANTES}
Nível do aluno: {NIVEL_DO_ALUNO}
Situação do conteúdo: {SITUACAO_CONTEUDO}
Prioridade: {PRIORIDADE_DO_CONTEUDO}
Modo de estudo: {MODO_ESTUDO}

IMPORTANTE
• usar as fontes apenas para entender a habilidade cobrada
• não copiar exemplos do livro
• criar exemplos inéditos
• manter linguagem clara, visual e respeitosa
• não infantilizar
""", values)

    if kind == "video":
        body = """MATERIAL PARA NOTEBOOKLM STUDIO — VIDEO OVERVIEW

OBJETIVO
Criar um Video Overview realmente visual, em português do Brasil, para uma criança de 9 anos, com foco em clareza, atenção curta e compreensão progressiva.

INSTRUÇÕES OBRIGATÓRIAS
• transformar o conteúdo em apresentação visual narrada
• organizar em sequência lógica, como mini aula
• usar exemplos inéditos
• priorizar elementos visuais e comparações concretas
• incluir 1 gancho inicial, 3 exemplos progressivos e 1 mini desafio final
• destacar 1 erro comum

ESTRUTURA
1. abertura com gancho
2. explicação visual da ideia central
3. exemplo básico
4. exemplo de aplicação
5. desafio leve
6. erro comum
7. fechamento com mini desafio
"""
    elif kind == "audio":
        body = """MATERIAL PARA NOTEBOOKLM STUDIO — AUDIO OVERVIEW

OBJETIVO
Criar um Audio Overview curto, dinâmico e claro, em português do Brasil, como se fosse um professor particular explicando o conteúdo ao aluno.

INSTRUÇÕES OBRIGATÓRIAS
• som de conversa guiada, não palestra longa
• frases curtas
• incluir perguntas para o aluno pensar
• incluir 3 exemplos progressivos
• reforçar 1 erro comum
• fechar com mini revisão e encorajamento
"""
    elif kind == "slides":
        body = """MATERIAL PARA NOTEBOOKLM STUDIO — SLIDES

OBJETIVO
Criar slides curtos, visuais e claros, para uma criança de 9 anos com atenção curta.

INSTRUÇÕES OBRIGATÓRIAS
• poucos slides
• cada slide com título curto
• no máximo 2 a 4 pontos por slide
• usar exemplos visuais e concretos
• progressão:
  - slide 1: ideia central
  - slide 2: exemplo básico
  - slide 3: aplicação
  - slide 4: erro comum
  - slide 5: mini desafio
"""
    elif kind == "quiz":
        body = """MATERIAL PARA NOTEBOOKLM STUDIO — QUIZ / FLASHCARDS

OBJETIVO
Criar um quiz rápido ou flashcards para revisão ativa, com progressão de dificuldade.

INSTRUÇÕES OBRIGATÓRIAS
• 5 itens no máximo
• começar com confiança
• avançar para aplicação
• terminar com desafio leve
• incluir pelo menos 1 erro comum
• respostas com explicação curtíssima
"""
    else:
        body = """MATERIAL PARA NOTEBOOKLM STUDIO — REPORT / RESUMO

OBJETIVO
Criar um resumo estratégico do conteúdo do dia, visualmente organizado, destacando o que mais importa.

INSTRUÇÕES OBRIGATÓRIAS
• resumir a ideia central
• listar pontos-chave
• destacar erro comum
• incluir 1 exemplo curto
• incluir 1 mini desafio
• manter texto enxuto e muito claro
"""
    return common + "\n\n" + body

def iter_paragraphs_in_table(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nested in cell.tables:
                for np in iter_paragraphs_in_table(nested):
                    yield np

def all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for p in iter_paragraphs_in_table(table):
            yield p
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for p in section.footer.paragraphs:
            yield p

def extract_placeholders(doc):
    found = set()
    for p in all_paragraphs(doc):
        found.update(PLACEHOLDER_RE.findall(p.text.replace("\n", "")))
    return sorted(found)

def get_base_block_text_from_doc(doc):
    paragraphs = list(doc.paragraphs)
    start = None
    for i, p in enumerate(paragraphs):
        if "0. BLOCO BASE GLOBAL" in p.text.upper():
            start = i
            break
    if start is None:
        return ""
    collected = []
    for p in paragraphs[start + 1:]:
        txt = p.text.strip()
        if re.match(r"^\d+\.", txt):
            break
        collected.append(p.text)
    return "\n".join(collected).strip()

def replace_in_runs(paragraph, replacements):
    full = paragraph.text
    new = full
    for key, value in replacements.items():
        pattern = r"\[\s*" + re.escape(key) + r"\s*\]"
        new = re.sub(pattern, value, new)
    if new != full:
        if paragraph.runs:
            paragraph.runs[0].text = new
            for r in paragraph.runs[1:]:
                r.text = ""
        else:
            paragraph.text = new

def replace_everywhere(doc, replacements):
    for p in all_paragraphs(doc):
        txt = p.text
        if BASE_MARKER in txt and replacements.get("__BASE_BLOCK__"):
            txt = txt.replace(BASE_MARKER, replacements["__BASE_BLOCK__"])
            if p.runs:
                p.runs[0].text = txt
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.text = txt
        replace_in_runs(p, replacements)

def generate_doc_from_template(uploaded_file, values):
    uploaded_file.seek(0)
    doc = Document(uploaded_file)
    vals = dict(values)
    vals["__BASE_BLOCK__"] = get_base_block_text_from_doc(doc) or safe_format(BASE_BLOCK_TEMPLATE, vals)
    replace_everywhere(doc, vals)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def export_json(values):
    return json.dumps(values, ensure_ascii=False, indent=2).encode("utf-8")

st.markdown("""
<style>
.block-container {padding-top: 1.2rem; padding-bottom: 2rem;}
.small-card {border: 1px solid rgba(49,51,63,0.15); border-radius: 16px; padding: 14px 16px; background: rgba(250,250,252,0.8);}
</style>
""", unsafe_allow_html=True)

st.title("Prompts Gu — Smart Web V4")
st.caption("A aba 2 gera prompts separados para a aba Studio do NotebookLM.")

with st.sidebar:
    st.header("Perfil base")
    for k, default in DEFAULTS.items():
        val = st.text_input(k, value=st.session_state.get(k, default), key="base_" + k)
        st.session_state[k] = val

tab1, tab2, tab3 = st.tabs(["1. Cronograma", "2. Studio", "3. DOCX"])

with tab1:
    st.subheader("Montar cronograma até a prova")
    a, b = st.columns(2)
    materia = a.text_input("Matéria", value=st.session_state.get("MATERIA", ""))
    data_hoje = a.text_input("Data de hoje", value=datetime.datetime.now().strftime("%d/%m/%Y"))
    data_prova = b.text_input("Data da prova", value=st.session_state.get("DATA_DA_PROVA", ""))
    outras = st.text_input("Outras provas no período", value=st.session_state.get("OUTRAS_PROVAS_NO_PERIODO", ""))
    conteudos = st.text_area("Conteúdos da prova", value=st.session_state.get("CONTEUDOS_DA_PROVA", ""), height=120)
    p1, p2, p3 = st.columns(3)
    prior = p1.text_area("Prioridade alta", value=st.session_state.get("PRIORIDADES", ""), height=120)
    medios = p2.text_area("Prioridade média", value=st.session_state.get("CONTEUDOS_MEDIOS", ""), height=120)
    sec = p3.text_area("Prioridade baixa", value=st.session_state.get("CONTEUDOS_SECUNDARIOS", ""), height=120)

    values = dict(st.session_state)
    values.update({"MATERIA": materia, "DATA_DE_HOJE": data_hoje, "DATA_DA_PROVA": data_prova, "OUTRAS_PROVAS_NO_PERIODO": outras, "CONTEUDOS_DA_PROVA": conteudos, "PRIORIDADES": prior, "CONTEUDOS_MEDIOS": medios, "CONTEUDOS_SECUNDARIOS": sec})
    values["BASE"] = safe_format(BASE_BLOCK_TEMPLATE, values)
    prompt_crono = safe_format(CRONOGRAMA_TEMPLATE, values)

    st.text_area("Prompt de cronograma", value=prompt_crono, height=420)
    c1, c2 = st.columns(2)
    c1.download_button("Baixar prompt .txt", prompt_crono.encode("utf-8"), file_name="prompt_cronograma_v4.txt")
    c2.download_button("Baixar JSON da aba", export_json(values), file_name="cronograma_valores_v4.json")

with tab2:
    st.subheader("Gerar instruções para a aba Studio do NotebookLM")
    c1, c2, c3 = st.columns(3)
    materia2 = c1.text_input("Matéria", value=st.session_state.get("MATERIA", ""), key="t2_materia")
    conteudo_dia = c1.text_input("Conteúdo do dia", value=st.session_state.get("CONTEUDO_DO_DIA", ""))
    data_hoje2 = c2.text_input("Data de hoje", value=datetime.datetime.now().strftime("%d/%m/%Y"), key="t2_hoje")
    data_prova2 = c2.text_input("Data da prova", value=st.session_state.get("DATA_DA_PROVA", ""), key="t2_prova")
    situacao = c3.selectbox("Situação do conteúdo", ["novo", "ja_visto", "em_dificuldade"], index=0)
    prioridade_conteudo = c3.selectbox("Prioridade", ["alta", "media", "baixa"], index=0)
    duracao = st.text_input("Duração desejada", value=st.session_state.get("DURACAO_DESEJADA", "15 a 25 minutos"))
    nivel = st.text_input("Nível do aluno", value=st.session_state.get("NIVEL_DO_ALUNO", DEFAULTS["NIVEL_DO_ALUNO"]))

    dias_restantes = "não calculado"
    try:
        hoje_dt = datetime.datetime.strptime(data_hoje2, "%d/%m/%Y")
        prova_dt = datetime.datetime.strptime(data_prova2, "%d/%m/%Y")
        days_num = (prova_dt - hoje_dt).days
        dias_restantes = str(days_num)
    except Exception:
        days_num = 5

    tipo = recommend_material(days_num, situacao, prioridade_conteudo)
    intensidade = recommend_intensity(days_num, situacao)
    modo = recommend_mode(days_num, situacao)

    k1, k2, k3 = st.columns(3)
    k1.markdown(f'<div class="small-card"><strong>Dias até a prova</strong><br>{dias_restantes}</div>', unsafe_allow_html=True)
    k2.markdown(f'<div class="small-card"><strong>Modo de estudo</strong><br>{modo}</div>', unsafe_allow_html=True)
    k3.markdown(f'<div class="small-card"><strong>Tipo recomendado</strong><br>{tipo}</div>', unsafe_allow_html=True)

    values2 = dict(st.session_state)
    values2.update({"MATERIA": materia2, "CONTEUDO_DO_DIA": conteudo_dia, "DATA_DE_HOJE": data_hoje2, "DATA_DA_PROVA": data_prova2, "DIAS_RESTANTES": dias_restantes, "SITUACAO_CONTEUDO": situacao, "PRIORIDADE_DO_CONTEUDO": prioridade_conteudo, "TIPO_DE_MATERIAL": tipo, "INTENSIDADE": intensidade, "MODO_ESTUDO": modo, "NIVEL_DO_ALUNO": nivel, "DURACAO_DESEJADA": duracao})
    values2["BASE"] = safe_format(BASE_BLOCK_TEMPLATE, values2)

    video_txt = make_prompt("video", values2)
    audio_txt = make_prompt("audio", values2)
    slides_txt = make_prompt("slides", values2)
    quiz_txt = make_prompt("quiz", values2)
    report_txt = make_prompt("report", values2)

    st.info("Use cada bloco abaixo no tipo certo de material dentro da aba Studio do NotebookLM.")

    with st.expander("Prompt para Video Overview", expanded=True):
        st.text_area("Video Overview", value=video_txt, height=260)
        st.download_button("Baixar prompt de vídeo", video_txt.encode("utf-8"), file_name="studio_video_overview_v4.txt")

    with st.expander("Prompt para Audio Overview"):
        st.text_area("Audio Overview", value=audio_txt, height=220)
        st.download_button("Baixar prompt de áudio", audio_txt.encode("utf-8"), file_name="studio_audio_overview_v4.txt")

    with st.expander("Prompt para Slides"):
        st.text_area("Slides", value=slides_txt, height=220)
        st.download_button("Baixar prompt de slides", slides_txt.encode("utf-8"), file_name="studio_slides_v4.txt")

    with st.expander("Prompt para Quiz / Flashcards"):
        st.text_area("Quiz / Flashcards", value=quiz_txt, height=220)
        st.download_button("Baixar prompt de quiz", quiz_txt.encode("utf-8"), file_name="studio_quiz_v4.txt")

    with st.expander("Prompt para Report / Resumo"):
        st.text_area("Report / Resumo", value=report_txt, height=200)
        st.download_button("Baixar prompt de report", report_txt.encode("utf-8"), file_name="studio_report_v4.txt")

    st.download_button("Baixar JSON da sessão", export_json(values2), file_name="studio_valores_v4.json")

with tab3:
    st.subheader("Preencher seu DOCX")
    uploaded = st.file_uploader("Envie o arquivo .docx", type=["docx"])
    if uploaded:
        uploaded.seek(0)
        doc_preview = Document(uploaded)
        placeholders = extract_placeholders(doc_preview)
        st.success(f"{len(placeholders)} campo(s) encontrado(s).")

        values3 = dict(st.session_state)
        c1, c2 = st.columns(2)
        values3["MATERIA"] = c1.text_input("MATERIA", value=values3.get("MATERIA", ""))
        values3["DATA_DE_HOJE"] = c1.text_input("DATA_DE_HOJE", value=datetime.datetime.now().strftime("%d/%m/%Y"))
        values3["DATA_DA_PROVA"] = c2.text_input("DATA_DA_PROVA", value=values3.get("DATA_DA_PROVA", ""))
        values3["CONTEUDO_DO_DIA"] = st.text_input("CONTEUDO_DO_DIA", value=values3.get("CONTEUDO_DO_DIA", ""))
        values3["CONTEUDOS_DA_PROVA"] = st.text_area("CONTEUDOS_DA_PROVA", value=values3.get("CONTEUDOS_DA_PROVA", ""), height=120)
        values3["OUTRAS_PROVAS_NO_PERIODO"] = st.text_input("OUTRAS_PROVAS_NO_PERIODO", value=values3.get("OUTRAS_PROVAS_NO_PERIODO", ""))
        p1, p2, p3 = st.columns(3)
        values3["PRIORIDADES"] = p1.text_area("PRIORIDADES", value=values3.get("PRIORIDADES", ""), height=100)
        values3["CONTEUDOS_MEDIOS"] = p2.text_area("CONTEUDOS_MEDIOS", value=values3.get("CONTEUDOS_MEDIOS", ""), height=100)
        values3["CONTEUDOS_SECUNDARIOS"] = p3.text_area("CONTEUDOS_SECUNDARIOS", value=values3.get("CONTEUDOS_SECUNDARIOS", ""), height=100)
        values3["DURACAO_DESEJADA"] = st.text_input("DURACAO_DESEJADA", value=values3.get("DURACAO_DESEJADA", "15 a 25 minutos"))

        try:
            hoje_dt = datetime.datetime.strptime(values3["DATA_DE_HOJE"], "%d/%m/%Y")
            prova_dt = datetime.datetime.strptime(values3["DATA_DA_PROVA"], "%d/%m/%Y") if values3["DATA_DA_PROVA"] else hoje_dt
            days_num = (prova_dt - hoje_dt).days
        except Exception:
            days_num = 5

        s1, s2 = st.columns(2)
        situ_doc = s1.selectbox("Situação do conteúdo do dia", ["novo", "ja_visto", "em_dificuldade"], index=0, key="doc_situ")
        pri_doc = s2.selectbox("Prioridade do conteúdo do dia", ["alta", "media", "baixa"], index=0, key="doc_pri")
        values3["TIPO_DE_MATERIAL"] = recommend_material(days_num, situ_doc, pri_doc)

        st.info("TIPO_DE_MATERIAL sugerido automaticamente: " + values3["TIPO_DE_MATERIAL"])

        with st.expander("Editar manualmente qualquer placeholder do DOCX"):
            for ph in placeholders:
                if ph not in values3:
                    values3[ph] = st.text_input(ph, value="", key="ph_" + ph)

        x1, x2 = st.columns(2)
        if x1.button("Preparar DOCX preenchido"):
            output = generate_doc_from_template(uploaded, values3)
            name = Path(uploaded.name).stem + "_preenchido_v4_" + datetime.datetime.now().strftime("%Y%m%d_%H%M") + ".docx"
            st.download_button("Baixar DOCX preenchido", data=output.getvalue(), file_name=name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        x2.download_button("Baixar JSON do DOCX", export_json(values3), file_name="docx_valores_v4.json")
    else:
        st.info("Envie o arquivo de prompts em .docx para preencher aqui.")
